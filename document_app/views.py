def index(request):
    return render(request, 'index.html')


import os
import io
from datetime import date
from django.conf import settings
from django.shortcuts import render
from django.http import HttpResponse
from .forms import DocumentForm
from .models import DocumentData
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches

def pdf_to_images(pdf_file, zoom_x=2, zoom_y=2):
    images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        # Set zoom factor to increase resolution
        matrix = fitz.Matrix(zoom_x, zoom_y)  # zoom_x, zoom_y control the DPI
        pix = page.get_pixmap(matrix=matrix)
        img = Image.open(io.BytesIO(pix.tobytes("png")))  # Convert pixmap to image
        images.append(img)
    return images

#this is the document generator for lower completion
def generate_document(
    data, images1, images2, images3, images4, images5, 
    images6, images7, images8, images9, images10, images11, image12, image13, image14, image15, image16, image17
):
    template_path = os.path.join(settings.BASE_DIR, 'document_app', 'templates', 'document_app', 'templates.docx')
    document = Document(template_path)

    # Replace form data placeholders with actual values
    for paragraph in document.paragraphs:
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    def insert_images_at_placeholder(placeholder, images, max_width=6.5, max_height=9):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        # Clear existing content in the cell
                        for paragraph in cell.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)
                        
                        # Insert images into the cell
                        for image in images:
                            image_stream = io.BytesIO()
                            image.save(image_stream, format='PNG')
                            image_stream.seek(0)
                            
                            new_paragraph = cell.add_paragraph()
                            run = new_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(max_width), height=Inches(max_height))  # Adjust sizes as needed

                        # Stop searching once we have inserted images
                        return

    # Insert images into cells with placeholders
    insert_images_at_placeholder('{{well_trajectory}}', images1)
    insert_images_at_placeholder('{{data}}', images2)
    insert_images_at_placeholder('{{schematic}}', images3)
    insert_images_at_placeholder('{{material_consumption}}', images4)
    insert_images_at_placeholder('{{tdas}}', images5)
    insert_images_at_placeholder('{{dsr}}', images6)
    insert_images_at_placeholder('{{mfiv_assembly}}', images7)
    insert_images_at_placeholder('{{quantum_packer}}', images8)
    insert_images_at_placeholder('{{washdown}}', images9)
    insert_images_at_placeholder('{{tallies}}', images10)
    insert_images_at_placeholder('{{csr}}', images11)

    # Insert the line_test image directly
    if image12:
        insert_images_at_placeholder('{{line_test}}', [image12])
    
    if image13:
        insert_images_at_placeholder('{{packer_setting}}', [image13])
    
    if image14:
        insert_images_at_placeholder('{{annulus_test}}', [image14])
    
    if image15:
        insert_images_at_placeholder('{{release_service}}', [image15])
    
    if image16:
        insert_images_at_placeholder('{{expand_ballseat}}', [image16])
    
    if image17:
        insert_images_at_placeholder('{{mfiv}}', [image17])

    # Save the modified document to a BytesIO object
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output

#this view is for lower completion for oil producers
def document_view(request):
    if request.method == 'POST':
        # Collect all non-media fields from request.POST as a dictionary
        data = request.POST.dict()

        print("Received data:", data)


        # Handle file uploads
        pdf_file1 = request.FILES.get('well_trajectory')
        pdf_file2 = request.FILES.get('data')
        pdf_file3 = request.FILES.get('schematic')
        pdf_file4 = request.FILES.get('material_consumption')
        pdf_file5 = request.FILES.get('tdas')
        pdf_file6 = request.FILES.get('dsr')
        pdf_file7 = request.FILES.get('quantum_packer')
        pdf_file8 = request.FILES.get('mfiv_assembly')
        pdf_file9 = request.FILES.get('washdown')
        pdf_file10 = request.FILES.get('csr')
        pdf_file11 = request.FILES.get('tallies')
        image_file12 = request.FILES.get('line_test')
        image_file13 = request.FILES.get('packer_setting')
        image_file14 = request.FILES.get('annulus_test')
        image_file15 = request.FILES.get('release_service')
        image_file16 = request.FILES.get('expand_ballseat')
        image_file17 = request.FILES.get('mfiv')
        

        # Convert PDFs to images if they exist
        images1 = pdf_to_images(pdf_file1) if pdf_file1 else []
        images2 = pdf_to_images(pdf_file2) if pdf_file2 else []
        images3 = pdf_to_images(pdf_file3) if pdf_file3 else []
        images4 = pdf_to_images(pdf_file4) if pdf_file4 else []
        images5 = pdf_to_images(pdf_file5) if pdf_file5 else []
        images6 = pdf_to_images(pdf_file6) if pdf_file6 else []
        images7 = pdf_to_images(pdf_file7) if pdf_file7 else []
        images8 = pdf_to_images(pdf_file8) if pdf_file8 else []
        images9 = pdf_to_images(pdf_file9) if pdf_file9 else []
        images10 = pdf_to_images(pdf_file10) if pdf_file10 else []
        images11 = pdf_to_images(pdf_file11) if pdf_file11 else []

         # Read the image directly without conversion if it exists
        image12 = Image.open(image_file12) if image_file12 else None
        image13 = Image.open(image_file13) if image_file13 else None
        image14 = Image.open(image_file14) if image_file14 else None
        image15 = Image.open(image_file15) if image_file15 else None
        image16 = Image.open(image_file16) if image_file16 else None
        image17 = Image.open(image_file17) if image_file17 else None
        

        # Generate the document with the provided data and images
        doc_file = generate_document(data, images1, images2, images3, images4, images5, images6, images7, images8, images9, images10, images11, image12, image13, image14, image15, image16, image17)

        # Create an HTTP response to return the generated document
        response = HttpResponse(doc_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{data.get("well_name", "document")}_report.docx"'
        
        # Return the response to trigger the download
        return response

    # For GET requests, simply render the HTML form
    return render(request, 'job-1.html')



def upper_oil(request):
    return render(request, 'index.html')

def both_oil(request):
    return render(request, 'index.html')


def upper_lower_fibre(request):
    return render(request, 'index.html')

def lower_suspension_fibre(request):
    return render(request, 'index.html')

def water_injector1_generate(
    data, images1, images2, images3, images4, images5, images6, images7, 
                                     images8, images9, images10, 
                                     images11, images12, images13, images14, images15, images16, 
                                      images18, images19, images20, images21, images22, 
                                     images23, image25, image26, image27, image28, image29, 
                                     image30, image31, image32, image33, image34, image35, image36
):
    template_path = os.path.join(settings.BASE_DIR, 'document_app', 'templates', 'document_app', 'WATERINJECTOR1.docx')
    document = Document(template_path)

    # Replace form data placeholders with actual values
    for paragraph in document.paragraphs:
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    def insert_images_at_placeholder(placeholder, images, max_width=6.5, max_height=9):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        # Clear existing content in the cell
                        for paragraph in cell.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)
                        
                        # Insert images into the cell
                        for image in images:
                            image_stream = io.BytesIO()
                            image.save(image_stream, format='PNG')
                            image_stream.seek(0)
                            
                            new_paragraph = cell.add_paragraph()
                            run = new_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(max_width), height=Inches(max_height))  # Adjust sizes as needed

                        # Stop searching once we have inserted images
                        return

    # Insert images into cells with placeholders
    insert_images_at_placeholder('{{well_trajectory}}', images1)
    insert_images_at_placeholder('{{data}}', images2)
    insert_images_at_placeholder('{{SCHEMATIC}}', images3)
    insert_images_at_placeholder('{{PnMLC}}', images4)
    insert_images_at_placeholder('{{TDAS}}', images5)
    insert_images_at_placeholder('{{DSRLC}}', images6)
    insert_images_at_placeholder('{{mfiv_assembly}}', images7)
    insert_images_at_placeholder('{{quantum_packer}}', images8)
    insert_images_at_placeholder('{{washdown}}', images9)
    insert_images_at_placeholder('{{csr}}', images10)
    insert_images_at_placeholder('{{tallies}}', images11)
    insert_images_at_placeholder('{{SURVEY}}', images12)
    insert_images_at_placeholder('{{PnMUC}}', images13)
    insert_images_at_placeholder('{{DSRIC}}', images14)
    insert_images_at_placeholder('{{TDR}}', images15)
    insert_images_at_placeholder('{{stplc}}', images16)
    # insert_images_at_placeholder('{{washdown}}', images17)
    insert_images_at_placeholder('{{lower_swivel_assembly}}', images18)
    insert_images_at_placeholder('{{nipple}}', images19)
    insert_images_at_placeholder('{{sgmpdf}}', images20)
    insert_images_at_placeholder('{{psa}}', images21)
    insert_images_at_placeholder('{{stca}}', images22)
    insert_images_at_placeholder('{{neocard}}', images23)
    # insert_images_at_placeholder('{{csr}}', images24)
    




    # Insert the line_test image directly
    if image25:
        insert_images_at_placeholder('{{line_test}}', [image25])
    
    if image26:
        insert_images_at_placeholder('{{packer_setting}}', [image26])
    
    if image27:
        insert_images_at_placeholder('{{annulus_test}}', [image27])
    
    if image28:
        insert_images_at_placeholder('{{release_service}}', [image28])
    
    if image29:
        insert_images_at_placeholder('{{blow_ballseat}}', [image29])
    
    if image30:
        insert_images_at_placeholder('{{mfiv}}', [image30])
    
    if image31:
        insert_images_at_placeholder('{{sgm}}', [image31])

    if image32:
        insert_images_at_placeholder('{{gch}}', [image32])

    if image33:
        insert_images_at_placeholder('{{ga}}', [image33])

    if image34:
        insert_images_at_placeholder('{{side_port}}', [image34])

    if image35:
        insert_images_at_placeholder('{{mfiv_opening}}', [image35])

    if image36:
        insert_images_at_placeholder('{{final_gauge}}', [image36])

    # Save the modified document to a BytesIO object
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output






def water_injector_1(request):
    if request.method == 'POST':
        # Collect all non-media fields from request.POST as a dictionary
        data = request.POST.dict()

        print("Received data:", data)


        # Handle file uploads
        pdf_file1 = request.FILES.get('well_trajectory')
        pdf_file2 = request.FILES.get('data')
        pdf_file3 = request.FILES.get('SCHEMATIC')
        pdf_file4 = request.FILES.get('PnMLC')
        pdf_file5 = request.FILES.get('TDAS')
        pdf_file6 = request.FILES.get('DSRLC')
        pdf_file7 = request.FILES.get('quantum_packer')
        pdf_file8 = request.FILES.get('mfiv_assembly')
        pdf_file9 = request.FILES.get('washdown')
        pdf_file10 = request.FILES.get('csr')
        pdf_file11 = request.FILES.get('tallies')
        pdf_file12 = request.FILES.get('SURVEY')
        pdf_file13 = request.FILES.get('PnMUC')
        pdf_file14 = request.FILES.get('DSRIC')
        pdf_file15 = request.FILES.get('TDR')
        pdf_file16 = request.FILES.get('stplc')
        # pdf_file17 = request.FILES.get('washdown')
        pdf_file18 = request.FILES.get('lower_swivel_assembly')
        pdf_file19 = request.FILES.get('nipple')
        pdf_file20 = request.FILES.get('sgmpdf')
        pdf_file21 = request.FILES.get('psa')
        pdf_file22 = request.FILES.get('stca')
        pdf_file23 = request.FILES.get('neocard')
        # pdf_file24 = request.FILES.get('csr')


        #getting the image files
        image_file1 = request.FILES.get('line_test')
        image_file2 = request.FILES.get('packer_setting')
        image_file3 = request.FILES.get('annulus_test')
        image_file4 = request.FILES.get('release_service')
        image_file5 = request.FILES.get('blow_ballseat')
        image_file6 = request.FILES.get('mfiv')
        image_file7 = request.FILES.get('sgm')
        image_file8 = request.FILES.get('gch')
        image_file9 = request.FILES.get('ga')
        image_file10 = request.FILES.get('side_port')
        image_file11 = request.FILES.get('mfiv_opening')
        image_file12 = request.FILES.get('final_gauge')
        

        

        # Convert PDFs to images if they exist
        images1 = pdf_to_images(pdf_file1) if pdf_file1 else []
        images2 = pdf_to_images(pdf_file2) if pdf_file2 else []
        images3 = pdf_to_images(pdf_file3) if pdf_file3 else []
        images4 = pdf_to_images(pdf_file4) if pdf_file4 else []
        images5 = pdf_to_images(pdf_file5) if pdf_file5 else []
        images6 = pdf_to_images(pdf_file6) if pdf_file6 else []
        images7 = pdf_to_images(pdf_file7) if pdf_file7 else []
        images8 = pdf_to_images(pdf_file8) if pdf_file8 else []
        images9 = pdf_to_images(pdf_file9) if pdf_file9 else []
        images10 = pdf_to_images(pdf_file10) if pdf_file10 else []
        images11 = pdf_to_images(pdf_file11) if pdf_file11 else []
        images12 = pdf_to_images(pdf_file12) if pdf_file12 else []
        images13 = pdf_to_images(pdf_file13) if pdf_file13 else []
        images14 = pdf_to_images(pdf_file14) if pdf_file14 else []
        images15 = pdf_to_images(pdf_file15) if pdf_file15 else []
        images16 = pdf_to_images(pdf_file16) if pdf_file16 else []
       
        images18 = pdf_to_images(pdf_file18) if pdf_file18 else []
        images19 = pdf_to_images(pdf_file19) if pdf_file19 else []
        images20 = pdf_to_images(pdf_file20) if pdf_file20 else []
        images21 = pdf_to_images(pdf_file21) if pdf_file21 else []
        images22 = pdf_to_images(pdf_file22) if pdf_file22 else []
        images23 = pdf_to_images(pdf_file23) if pdf_file23 else []
        # images24 = pdf_to_images(pdf_file24) if pdf_file24 else []
        

         # Read the image directly without conversion if it exists
        image25 = Image.open(image_file1) if image_file1 else None
        image26 = Image.open(image_file2) if image_file2 else None
        image27 = Image.open(image_file3) if image_file3 else None
        image28 = Image.open(image_file4) if image_file4 else None
        image29 = Image.open(image_file5) if image_file5 else None
        image30 = Image.open(image_file6) if image_file6 else None
        image31 = Image.open(image_file7) if image_file7 else None
        image32 = Image.open(image_file8) if image_file8 else None
        image33 = Image.open(image_file9) if image_file9 else None
        image34 = Image.open(image_file10) if image_file10 else None
        image35 = Image.open(image_file11) if image_file11 else None
        image36 = Image.open(image_file12) if image_file12 else None
        

        # Generate the document with the provided data and images
        doc_file = water_injector1_generate(data, images1, images2, images3, images4, images5, images6, images7, 
                                     images8, images9, images10, 
                                     images11, images12, images13, images14, images15, images16, 
                                      images18, images19, images20, images21, images22, 
                                     images23, image25, image26, image27, image28, image29, 
                                     image30, image31, image32, image33, image34, image35, image36)

        # Create an HTTP response to return the generated document
        response = HttpResponse(doc_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{data.get("well_name", "document")}_report.docx"'

        return response
        
        # Return the response to trigger the download
    return render(request, 'water_injector1.html')










def water_injector2_generate(
    data, images1, images2, images3, images4, images5, images6, images7, 
                                     images8, images9, images10, 
                                     images11, images12, images13, images14, image15, image16, image17,
                                      image18, image19, image20, image21, image22
                                    
):
    template_path = os.path.join(settings.BASE_DIR, 'document_app', 'templates', 'document_app', 'WATERINJECTOR2.docx')
    document = Document(template_path)

    # Replace form data placeholders with actual values
    for paragraph in document.paragraphs:
        for key, value in data.items():
            placeholder = '{{' + key + '}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = '{{' + key + '}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    def insert_images_at_placeholder(placeholder, images, max_width=6.5, max_height=9):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        # Clear existing content in the cell
                        for paragraph in cell.paragraphs:
                            p = paragraph._element
                            p.getparent().remove(p)
                        
                        # Insert images into the cell
                        for image in images:
                            image_stream = io.BytesIO()
                            image.save(image_stream, format='PNG')
                            image_stream.seek(0)
                            
                            new_paragraph = cell.add_paragraph()
                            run = new_paragraph.add_run()
                            run.add_picture(image_stream, width=Inches(max_width), height=Inches(max_height))  # Adjust sizes as needed

                        # Stop searching once we have inserted images
                        return

    # Insert images into cells with placeholders
    insert_images_at_placeholder('{{well_trajectory}}', images1)
    insert_images_at_placeholder('{{SCHEMATIC}}', images2)
    insert_images_at_placeholder('{{PnMLC}}', images3)
    insert_images_at_placeholder('{{TDAS}}', images4)
    insert_images_at_placeholder('{{jorunals}}', images5)
    insert_images_at_placeholder('{{ic_quantum}}', images6)
    insert_images_at_placeholder('{{lower_quantum_packer}}', images7)
    insert_images_at_placeholder('{{washdown}}', images8)
    insert_images_at_placeholder('{{csr}}', images9)
    insert_images_at_placeholder('{{tallies_lc}}', images10)
    insert_images_at_placeholder('{{SURVEY}}', images11)
    insert_images_at_placeholder('{{nipple}}', images12)
    insert_images_at_placeholder('{{tallies_ic}}', images13)
    insert_images_at_placeholder('{{weg}}', images14)
   
    




    # Insert the line_test image directly
    if image15:
        insert_images_at_placeholder('{{line_test}}', [image15])
    
    if image16:
        insert_images_at_placeholder('{{packer_setting}}', [image16])
    
    if image17:
        insert_images_at_placeholder('{{annulus_test}}', [image17])
    
    if image18:
        insert_images_at_placeholder('{{release_service}}', [image18])
    
    if image19:
        insert_images_at_placeholder('{{blow_ballseat}}', [image19])
    
    if image20:
        insert_images_at_placeholder('{{lbfv}}', [image20])
    
    if image21:
        insert_images_at_placeholder('{{i_line_test}}', [image21])

    if image22:
        insert_images_at_placeholder('{{gch}}', [image22])

    
    # Save the modified document to a BytesIO object
    output = io.BytesIO()
    document.save(output)
    output.seek(0)

    return output






def water_injector_2(request):
    if request.method == 'POST':
        # Collect all non-media fields from request.POST as a dictionary
        data = request.POST.dict()

        print("Received data:", data)


        # Handle file uploads
        pdf_file1 = request.FILES.get('well_trajectory')
        
        pdf_file2 = request.FILES.get('SCHEMATIC')
        pdf_file3 = request.FILES.get('PnMLC')
        pdf_file4 = request.FILES.get('TDAS')
        pdf_file5 = request.FILES.get('jorunals')
        pdf_file6 = request.FILES.get('ic_quantum')
        pdf_file7 = request.FILES.get('lower_quantum_packer')
        pdf_file8 = request.FILES.get('washdown')
        pdf_file9 = request.FILES.get('csr')
        pdf_file10 = request.FILES.get('tallies_lc')
        pdf_file11 = request.FILES.get('SURVEY')
        
        pdf_file12 = request.FILES.get('nipple')
        
        pdf_file13 = request.FILES.get('tallies_ic')
        pdf_file14 = request.FILES.get('weg')
        


        #getting the image files
        image_file1 = request.FILES.get('line_test')
        image_file2 = request.FILES.get('packer_setting')
        image_file3 = request.FILES.get('annulus_test')
        image_file4 = request.FILES.get('release_service')
        image_file5 = request.FILES.get('blow_ballseat')
        image_file6 = request.FILES.get('lbfv')
        image_file7 = request.FILES.get('i_line_test')
        image_file8 = request.FILES.get('inflow_test')
        
        

        

        # Convert PDFs to images if they exist
        images1 = pdf_to_images(pdf_file1) if pdf_file1 else []
        images2 = pdf_to_images(pdf_file2) if pdf_file2 else []
        images3 = pdf_to_images(pdf_file3) if pdf_file3 else []
        images4 = pdf_to_images(pdf_file4) if pdf_file4 else []
        images5 = pdf_to_images(pdf_file5) if pdf_file5 else []
        images6 = pdf_to_images(pdf_file6) if pdf_file6 else []
        images7 = pdf_to_images(pdf_file7) if pdf_file7 else []
        images8 = pdf_to_images(pdf_file8) if pdf_file8 else []
        images9 = pdf_to_images(pdf_file9) if pdf_file9 else []
        images10 = pdf_to_images(pdf_file10) if pdf_file10 else []
        images11 = pdf_to_images(pdf_file11) if pdf_file11 else []
        
       
        images12 = pdf_to_images(pdf_file12) if pdf_file12 else []
        
        images13 = pdf_to_images(pdf_file13) if pdf_file13 else []
        images14 = pdf_to_images(pdf_file14) if pdf_file14 else []
        # images24 = pdf_to_images(pdf_file24) if pdf_file24 else []
        

         # Read the image directly without conversion if it exists
        image15 = Image.open(image_file1) if image_file1 else None
        image16 = Image.open(image_file2) if image_file2 else None
        image17 = Image.open(image_file3) if image_file3 else None
        image18 = Image.open(image_file4) if image_file4 else None
        image19 = Image.open(image_file5) if image_file5 else None
        image20 = Image.open(image_file6) if image_file6 else None
        image21 = Image.open(image_file7) if image_file7 else None
        image22 = Image.open(image_file8) if image_file8 else None
       
        

        # Generate the document with the provided data and images
        doc_file = water_injector2_generate(data, images1, images2, images3, images4, images5, images6, images7, 
                                     images8, images9, images10, 
                                     images11, images12, images13, images14, image15, image16, 
                                      image17, image18, image19, image20, image21, image22, 
                                    )

        # Create an HTTP response to return the generated document
        response = HttpResponse(doc_file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{data.get("well_name", "document")}_report.docx"'

        return response
        
        # Return the response to trigger the download
    return render(request, 'water_injector2.html')

